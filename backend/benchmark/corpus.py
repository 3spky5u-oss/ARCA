"""
Corpus Converter — .docx to Markdown for benchmark ingestion.

Uses python-docx (already in requirements) to extract:
- Heading styles → markdown headers (#, ##, ###)
- Tables → markdown tables
- Bold/italic → **bold**/*italic*
- Paragraph text preserved as-is
"""
import logging
import re
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


class DocxConverter:
    """Convert .docx files to clean Markdown."""

    def convert_file(self, docx_path: str, output_dir: str) -> Optional[str]:
        """Convert a single .docx to .md.

        Args:
            docx_path: Path to .docx file
            output_dir: Directory for .md output

        Returns:
            Path to output .md file, or None on failure
        """
        from docx import Document

        docx_path = Path(docx_path)
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        output_path = output_dir / f"{docx_path.stem}.md"

        try:
            doc = Document(str(docx_path))
            md_lines = []

            for element in doc.element.body:
                tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

                if tag == 'p':
                    # Paragraph
                    para = None
                    for p in doc.paragraphs:
                        if p._element is element:
                            para = p
                            break
                    if para is None:
                        continue

                    line = self._convert_paragraph(para)
                    if line is not None:
                        md_lines.append(line)

                elif tag == 'tbl':
                    # Table
                    tbl = None
                    for t in doc.tables:
                        if t._element is element:
                            tbl = t
                            break
                    if tbl is not None:
                        table_md = self._convert_table(tbl)
                        md_lines.append(table_md)

            md_text = "\n\n".join(md_lines)

            # Clean up excessive blank lines
            md_text = re.sub(r'\n{3,}', '\n\n', md_text)

            output_path.write_text(md_text, encoding="utf-8")
            logger.info(f"Converted {docx_path.name} -> {output_path.name} ({len(md_text)} chars)")
            return str(output_path)

        except Exception as e:
            logger.error(f"Failed to convert {docx_path.name}: {e}")
            return None

    def convert_directory(self, corpus_dir: str, output_dir: str) -> List[str]:
        """Convert all .docx files in a directory.

        Returns list of output .md file paths.
        """
        corpus_dir = Path(corpus_dir)
        output_dir = Path(output_dir)

        docx_files = sorted(corpus_dir.glob("*.docx"))
        if not docx_files:
            logger.warning(f"No .docx files found in {corpus_dir}")
            return []

        logger.info(f"Converting {len(docx_files)} .docx files to markdown")

        results = []
        for docx_file in docx_files:
            result = self.convert_file(str(docx_file), str(output_dir))
            if result:
                results.append(result)

        logger.info(f"Successfully converted {len(results)}/{len(docx_files)} files")
        return results

    def _convert_paragraph(self, para) -> Optional[str]:
        """Convert a python-docx paragraph to markdown."""
        style_name = (para.style.name or "").lower() if para.style else ""
        text = para.text.strip()

        if not text:
            return None

        # Heading detection
        if "heading 1" in style_name or style_name == "title":
            return f"# {text}"
        elif "heading 2" in style_name:
            return f"## {text}"
        elif "heading 3" in style_name:
            return f"### {text}"
        elif "heading 4" in style_name:
            return f"#### {text}"
        elif "heading" in style_name:
            # Generic heading
            return f"### {text}"

        # Process inline formatting (bold/italic)
        formatted = self._format_runs(para.runs)
        return formatted if formatted.strip() else None

    def _format_runs(self, runs) -> str:
        """Convert paragraph runs to markdown with bold/italic."""
        parts = []
        for run in runs:
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            parts.append(text)
        return "".join(parts)

    def _convert_table(self, table) -> str:
        """Convert a python-docx table to markdown table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if not rows:
            return ""

        # Add separator after header row
        if len(rows) >= 1:
            n_cols = len(table.rows[0].cells)
            separator = "| " + " | ".join(["---"] * n_cols) + " |"
            rows.insert(1, separator)

        return "\n".join(rows)
