"""
Cohesionn Session Knowledge - Temporary per-session document indexing

Allows users to upload documents (RFPs, specs, reports) that get chunked
and indexed for the duration of their session.

Uses in-memory Qdrant for truly ephemeral storage - no disk usage,
automatic cleanup when session object is garbage collected.

Usage:
    from tools.cohesionn.session import SessionKnowledge

    # Create session store
    session = SessionKnowledge(session_id="user123")

    # Add a document
    result = session.add_document("/path/to/RFP.pdf", "rfp")

    # Search session documents
    results = session.search("what are the deliverables")

    # Clean up when session ends
    session.clear()
"""

import logging
import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    VectorParams,
    PointStruct,
    Filter,
    FieldCondition,
    MatchValue,
    MatchAny,
)
import hashlib
from datetime import datetime, timedelta

from .embeddings import get_embedder
from .chunker import SemanticChunker, Chunk
from .reranker import get_reranker

logger = logging.getLogger(__name__)


def chunk_id_to_int(chunk_id: str) -> int:
    """Convert hex chunk_id to integer for Qdrant point ID."""
    return int(chunk_id, 16)

# Vector dimension for Qwen3-Embedding-0.6B
VECTOR_DIM = 1024


@dataclass
class SessionDocument:
    """A document added to the session"""

    doc_id: str
    filename: str
    doc_type: str  # "rfp", "spec", "report", "other"
    chunk_count: int
    added_at: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class SessionSearchResult:
    """Result from session knowledge search"""

    content: str
    score: float
    doc_id: str
    filename: str
    doc_type: str
    page: Optional[int]
    section: Optional[str]


class SessionKnowledge:
    """
    Temporary in-memory knowledge store for user session.

    Documents uploaded during a session are chunked, embedded, and stored
    in an ephemeral in-memory Qdrant instance. When the session ends,
    the data is automatically cleaned up.
    """

    def __init__(self, session_id: str, persist: bool = False):
        """
        Args:
            session_id: Unique session identifier
            persist: If True, persist to disk (default: in-memory only)
        """
        self.session_id = session_id
        self.collection_name = f"session_{session_id}"

        # Session timestamps for cleanup
        self.created_at = datetime.now()
        self.last_accessed = datetime.now()

        # Use in-memory Qdrant client by default
        if persist:
            self.temp_dir = Path(tempfile.mkdtemp(prefix="cohesionn_session_"))
            self.client = QdrantClient(path=str(self.temp_dir))
        else:
            # Truly ephemeral - in-memory only
            self.client = QdrantClient(":memory:")
            self.temp_dir = None

        # Create session collection
        self.client.create_collection(
            collection_name=self.collection_name,
            vectors_config=VectorParams(
                size=VECTOR_DIM,
                distance=Distance.COSINE,
            ),
        )

        self.embedder = get_embedder()
        self.reranker = get_reranker()
        # Larger chunks = better context, smaller overlap = more chunks
        # 96GB RAM can handle it
        self.chunker = SemanticChunker(chunk_size=2500, chunk_overlap=0)

        # Track documents in session
        self.documents: Dict[str, SessionDocument] = {}

        logger.info(f"SessionKnowledge initialized: {session_id}")

    def add_document(
        self,
        file_path: Path,
        doc_type: str = "other",
        metadata: Dict[str, Any] = None,
    ) -> SessionDocument:
        """
        Add a document to the session knowledge.

        Args:
            file_path: Path to document (PDF, DOCX, TXT)
            doc_type: Type hint ("rfp", "spec", "report", "other")
            metadata: Additional metadata

        Returns:
            SessionDocument with info about the added document
        """
        file_path = Path(file_path)
        metadata = metadata or {}

        # Generate doc ID
        doc_id = hashlib.md5(f"{self.session_id}:{file_path.name}:{datetime.now().isoformat()}".encode()).hexdigest()[
            :10
        ]

        logger.info(f"Adding document: {file_path.name} as {doc_type}")

        # Extract text
        text, pages = self._extract_text(file_path)

        if not text:
            raise ValueError(f"Could not extract text from {file_path}")

        # Build chunk metadata
        base_metadata = {
            "doc_id": doc_id,
            "filename": file_path.name,
            "doc_type": doc_type,
            "session_id": self.session_id,
            **metadata,
        }

        # Chunk the document
        if pages:
            chunks = self.chunker.chunk_pages(pages, base_metadata)
        else:
            chunks = self.chunker.chunk_text(text, base_metadata)

        if not chunks:
            raise ValueError(f"No chunks created from {file_path}")

        # Embed and store
        self._add_chunks(chunks)

        # Track document
        session_doc = SessionDocument(
            doc_id=doc_id,
            filename=file_path.name,
            doc_type=doc_type,
            chunk_count=len(chunks),
            added_at=datetime.now().isoformat(),
            metadata=metadata,
        )
        self.documents[doc_id] = session_doc

        logger.info(f"Added {len(chunks)} chunks from {file_path.name}")

        return session_doc

    def add_text(
        self,
        text: str,
        name: str = "pasted_text",
        doc_type: str = "other",
        metadata: Dict[str, Any] = None,
    ) -> SessionDocument:
        """
        Add raw text to session knowledge.

        Args:
            text: Text content to add
            name: Name for the text block
            doc_type: Type hint
            metadata: Additional metadata

        Returns:
            SessionDocument
        """
        metadata = metadata or {}

        doc_id = hashlib.md5(f"{self.session_id}:{name}:{datetime.now().isoformat()}".encode()).hexdigest()[:10]

        base_metadata = {
            "doc_id": doc_id,
            "filename": name,
            "doc_type": doc_type,
            "session_id": self.session_id,
            **metadata,
        }

        chunks = self.chunker.chunk_text(text, base_metadata)

        if not chunks:
            raise ValueError("No chunks created from text")

        self._add_chunks(chunks)

        session_doc = SessionDocument(
            doc_id=doc_id,
            filename=name,
            doc_type=doc_type,
            chunk_count=len(chunks),
            added_at=datetime.now().isoformat(),
            metadata=metadata,
        )
        self.documents[doc_id] = session_doc

        return session_doc

    def search(
        self,
        query: str,
        top_k: int = 5,
        doc_types: Optional[List[str]] = None,
        rerank: bool = True,
    ) -> List[SessionSearchResult]:
        """
        Search session knowledge.

        Args:
            query: Search query
            top_k: Number of results
            doc_types: Filter by document types
            rerank: Whether to apply cross-encoder reranking

        Returns:
            List of SessionSearchResult
        """
        self.touch()  # Update last_accessed on search

        if self.chunk_count == 0:
            return []

        # Build filter
        filter_conditions = None
        if doc_types:
            filter_conditions = Filter(
                must=[
                    FieldCondition(key="doc_type", match=MatchAny(any=doc_types))
                ]
            )

        # Embed query
        query_embedding = self.embedder.embed_query(query)

        # Search - with 96GB RAM, retrieve more candidates for better reranking
        initial_k = min(50, self.chunk_count)
        response = self.client.query_points(
            collection_name=self.collection_name,
            query=query_embedding,
            limit=initial_k,
            query_filter=filter_conditions,
            with_payload=True,
        )

        if not response.points:
            return []

        # Format results
        candidates = []
        for hit in response.points:
            payload = hit.payload or {}
            candidates.append({
                "content": payload.get("content", ""),
                "metadata": {k: v for k, v in payload.items() if k != "content"},
                "distance": 1 - hit.score,
                "score": hit.score,  # Qdrant returns similarity directly
            })

        # Rerank
        if rerank and len(candidates) > top_k:
            candidates = self.reranker.rerank(query, candidates, top_k=top_k)
        else:
            candidates = candidates[:top_k]

        # Build results
        search_results = []
        for c in candidates:
            meta = c["metadata"]
            search_results.append(
                SessionSearchResult(
                    content=c["content"],
                    score=c.get("rerank_score", c["score"]),
                    doc_id=meta.get("doc_id", ""),
                    filename=meta.get("filename", ""),
                    doc_type=meta.get("doc_type", "other"),
                    page=meta.get("page"),
                    section=meta.get("section"),
                )
            )

        return search_results

    def get_context(self, query: str, max_chunks: int = 5) -> str:
        """
        Get formatted context for LLM consumption.

        Args:
            query: Search query
            max_chunks: Maximum chunks to include

        Returns:
            Formatted context string
        """
        results = self.search(query, top_k=max_chunks)

        if not results:
            return ""

        parts = []
        for i, r in enumerate(results, 1):
            source = f"[{r.filename}"
            if r.page:
                source += f", p.{r.page}"
            source += "]"

            parts.append(f"--- {source} ---\n{r.content}")

        return "\n\n".join(parts)

    def remove_document(self, doc_id: str):
        """Remove a document from session knowledge"""
        if doc_id not in self.documents:
            return

        # Delete from collection using filter
        self.client.delete(
            collection_name=self.collection_name,
            points_selector=Filter(
                must=[
                    FieldCondition(key="doc_id", match=MatchValue(value=doc_id))
                ]
            ),
        )

        # Remove from tracking
        del self.documents[doc_id]

        logger.info(f"Removed document: {doc_id}")

    def clear(self):
        """Clear all session knowledge"""
        # Delete and recreate collection
        try:
            self.client.delete_collection(self.collection_name)
        except Exception:
            pass

        # Recreate empty collection
        self.client.create_collection(
            collection_name=self.collection_name,
            vectors_config=VectorParams(
                size=VECTOR_DIM,
                distance=Distance.COSINE,
            ),
        )

        # Clear tracking
        self.documents.clear()

        logger.info(f"Cleared session: {self.session_id}")

    def cleanup(self):
        """Full cleanup - call when session ends"""
        self.clear()

        # Remove temp directory if persisted
        if self.temp_dir and self.temp_dir.exists():
            shutil.rmtree(self.temp_dir, ignore_errors=True)

        logger.info(f"Session cleanup complete: {self.session_id}")

    def list_documents(self) -> List[SessionDocument]:
        """List all documents in session"""
        return list(self.documents.values())

    @property
    def chunk_count(self) -> int:
        """Total chunks in session"""
        try:
            result = self.client.count(collection_name=self.collection_name)
            return result.count
        except Exception:
            return 0

    @property
    def document_count(self) -> int:
        """Number of documents in session"""
        return len(self.documents)

    def touch(self):
        """Update last_accessed timestamp"""
        self.last_accessed = datetime.now()

    def is_stale(self, max_age_hours: int = 24) -> bool:
        """Check if session has been inactive longer than max_age_hours"""
        cutoff = datetime.now() - timedelta(hours=max_age_hours)
        return self.last_accessed < cutoff

    def _add_chunks(self, chunks: List[Chunk]):
        """Add chunks to collection"""
        if not chunks:
            return

        documents = [c.content for c in chunks]

        # Embed
        embeddings = self.embedder.embed_documents(documents)

        # Build points - convert hex chunk_id to integer for Qdrant
        points = [
            PointStruct(
                id=chunk_id_to_int(c.chunk_id),
                vector=emb,
                payload={
                    "content": c.content,
                    "chunk_id": c.chunk_id,  # Store original for reference
                    **c.metadata,
                },
            )
            for c, emb in zip(chunks, embeddings)
        ]

        # Upsert
        self.client.upsert(
            collection_name=self.collection_name,
            points=points,
            wait=True,
        )

    def _extract_text(self, file_path: Path):
        """Extract text from file"""
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            import fitz

            doc = fitz.open(str(file_path))
            pages = []
            for i, page in enumerate(doc, 1):
                text = page.get_text("text")
                pages.append({"page_num": i, "text": text})
            doc.close()
            full_text = "\n\n".join(p["text"] for p in pages)
            return full_text, pages

        elif suffix == ".docx":
            from docx import Document

            doc = Document(str(file_path))
            parts = []
            for para in doc.paragraphs:
                parts.append(para.text)
            return "\n\n".join(parts), None

        elif suffix in {".txt", ".md"}:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return text, None

        else:
            raise ValueError(f"Unsupported file type: {suffix}")


# Session registry for managing multiple sessions
_sessions: Dict[str, SessionKnowledge] = {}


def get_session(session_id: str) -> SessionKnowledge:
    """Get or create a session knowledge store"""
    if session_id not in _sessions:
        _sessions[session_id] = SessionKnowledge(session_id)
    else:
        _sessions[session_id].touch()  # Update last_accessed on retrieval
    return _sessions[session_id]


def clear_session(session_id: str):
    """Clear and remove a session"""
    if session_id in _sessions:
        _sessions[session_id].cleanup()
        del _sessions[session_id]


def list_sessions() -> List[str]:
    """List active sessions"""
    return list(_sessions.keys())


def cleanup_stale_sessions(max_age_hours: int = 24) -> Dict[str, Any]:
    """
    Clean up sessions that have been inactive longer than max_age_hours.

    Args:
        max_age_hours: Maximum session age in hours before cleanup

    Returns:
        Dict with cleanup statistics
    """
    stale_sessions = []
    for session_id, session in list(_sessions.items()):
        if session.is_stale(max_age_hours):
            stale_sessions.append(session_id)

    cleaned = 0
    total_docs = 0
    total_chunks = 0

    for session_id in stale_sessions:
        session = _sessions[session_id]
        total_docs += session.document_count
        total_chunks += session.chunk_count
        session.cleanup()
        del _sessions[session_id]
        cleaned += 1
        logger.info(f"Cleaned up stale session: {session_id}")

    if cleaned > 0:
        logger.info(f"Session cleanup: {cleaned} sessions, {total_docs} docs, {total_chunks} chunks")

    return {
        "sessions_cleaned": cleaned,
        "documents_removed": total_docs,
        "chunks_removed": total_chunks,
    }


def get_session_stats() -> Dict[str, Any]:
    """Get statistics about active sessions"""
    return {
        "active_sessions": len(_sessions),
        "total_documents": sum(s.document_count for s in _sessions.values()),
        "total_chunks": sum(s.chunk_count for s in _sessions.values()),
        "sessions": {
            session_id: {
                "documents": s.document_count,
                "chunks": s.chunk_count,
                "created_at": s.created_at.isoformat(),
                "last_accessed": s.last_accessed.isoformat(),
            }
            for session_id, s in _sessions.items()
        },
    }
