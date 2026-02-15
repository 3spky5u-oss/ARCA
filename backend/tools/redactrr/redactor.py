"""
Redactrr - Agentic Document Redaction

Intelligently identifies and removes PII from documents for external analysis.
Handles Word (.docx) and digital PDF files.
"""

import logging
import re
import json
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from dataclasses import dataclass, field
from enum import Enum

from config import runtime_config
from domain_loader import get_pipeline_config

logger = logging.getLogger(__name__)


class PIIType(Enum):
    """Types of PII to redact"""

    PERSON_NAME = "person_name"
    COMPANY_NAME = "company_name"
    ADDRESS = "address"
    PHONE = "phone"
    EMAIL = "email"
    PROJECT_NUMBER = "project_number"
    FILE_NUMBER = "file_number"
    DATE = "date"
    LOCATION = "location"
    COORDINATES = "coordinates"
    LICENSE_PLATE = "license_plate"
    SIN_SSN = "sin_ssn"
    FINANCIAL = "financial"
    CUSTOM = "custom"


@dataclass
class PIIEntity:
    """A detected PII entity"""

    text: str
    pii_type: PIIType
    start: int = -1
    end: int = -1
    replacement: str = ""
    confidence: float = 1.0

    def __post_init__(self):
        if not self.replacement:
            self.replacement = self._generate_replacement()

    def _generate_replacement(self) -> str:
        """Generate appropriate redaction placeholder"""
        placeholders = {
            PIIType.PERSON_NAME: "[PERSON]",
            PIIType.COMPANY_NAME: "[COMPANY]",
            PIIType.ADDRESS: "[ADDRESS]",
            PIIType.PHONE: "[PHONE]",
            PIIType.EMAIL: "[EMAIL]",
            PIIType.PROJECT_NUMBER: "[PROJECT #]",
            PIIType.FILE_NUMBER: "[FILE #]",
            PIIType.DATE: "[DATE]",
            PIIType.LOCATION: "[LOCATION]",
            PIIType.COORDINATES: "[COORDINATES]",
            PIIType.LICENSE_PLATE: "[PLATE]",
            PIIType.SIN_SSN: "[ID #]",
            PIIType.FINANCIAL: "[FINANCIAL]",
            PIIType.CUSTOM: "[REDACTED]",
        }
        return placeholders.get(self.pii_type, "[REDACTED]")


@dataclass
class RedactionResult:
    """Result from redaction process"""

    original_path: str
    redacted_path: str
    success: bool
    entities_found: int
    entities_redacted: int
    entity_types: Dict[str, int]
    warnings: List[str] = field(default_factory=list)
    error: Optional[str] = None


class PIIDetector:
    """
    Detect PII using regex patterns and LLM analysis.

    Three-phase approach:
    1. Fast regex scan for obvious patterns (email, phone, etc.)
    2. Heuristic name detection (capitalized words in context)
    3. LLM analysis with smart sampling (header + footer, single pass)

    The LLM phase samples strategically from document header (first 3000 chars)
    and footer (last 1500 chars) where most PII appears, reducing processing
    time from 5+ minutes to ~90 seconds.
    """

    # Regex patterns for common PII
    PATTERNS = {
        PIIType.EMAIL: r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
        PIIType.PHONE: r"\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b",
        PIIType.COORDINATES: r"\b-?\d{1,3}\.\d{4,}[,\s]+-?\d{1,3}\.\d{4,}\b",
        PIIType.SIN_SSN: r"\b\d{3}[-\s]?\d{2}[-\s]?\d{4}\b",
        PIIType.LICENSE_PLATE: r"\b[A-Z]{2,3}[-\s]?[0-9]{2,4}[-\s]?[A-Z]{0,3}\b",
        PIIType.PROJECT_NUMBER: r"\b(?:Project|Proj\.?|File|Job|Contract|Work\s*Order)[\s#:\-]*[A-Z0-9][\w\-]{2,20}\b",
        PIIType.FILE_NUMBER: r"\b(?:File|Ref|Reference|Invoice|PO)[\s#:\-]*[A-Z0-9][\w\-]{2,15}\b",
    }

    # Common words that look like names but aren't
    _BASE_NAME_STOPWORDS = {
        # Common words
        "the",
        "and",
        "for",
        "with",
        "from",
        "this",
        "that",
        "will",
        "have",
        "been",
        "were",
        "are",
        "was",
        "has",
        "had",
        "may",
        "can",
        "per",
        "all",
        "any",
        "not",
        # Generic terms often capitalized in technical documents
        "tier",
        "phase",
        "stage",
        "section",
        "table",
        "figure",
        "appendix",
        "report",
        "site",
        "area",
        "zone",
        "level",
        "sample",
        "test",
        "analysis",
        "method",
        "data",
        "results",
        "standard",
        "specification",
        "assessment",
        "investigation",
        "review",
        "preliminary",
        "final",
        "draft",
        "revised",
        "updated",
        "approved",
        "fine",
        "coarse",
        "medium",
        "high",
        "low",
        "north",
        "south",
        "east",
        "west",
        "monday",
        "tuesday",
        "wednesday",
        "thursday",
        "friday",
        "saturday",
        "sunday",
        "january",
        "february",
        "march",
        "april",
        "june",
        "july",
        "august",
        "september",
        "october",
        "november",
        "december",
        "province",
        "city",
        "town",
        "county",
        "municipal",
        # Technical abbreviations
        "ltd",
        "inc",
        "corp",
        "llc",
        "eng",
        "msc",
        "bsc",
        "phd",
        "ing",
    }

    def __init__(self, use_llm: bool = True, model: str = None):
        self.use_llm = use_llm
        self.model = model or runtime_config.model_chat
        self._client = None
        # Build stopwords from base set + domain pipeline config
        domain_terms = get_pipeline_config().get("redactrr_technical_terms", set())
        self.NAME_STOPWORDS = set(self._BASE_NAME_STOPWORDS)
        self.NAME_STOPWORDS.update(t.lower() for t in domain_terms)

    @property
    def client(self):
        if self._client is None and self.use_llm:
            from utils.llm import get_llm_client

            self._client = get_llm_client("chat")
        return self._client

    def detect(self, text: str) -> List[PIIEntity]:
        """
        Detect all PII in text.

        Args:
            text: Document text to scan

        Returns:
            List of detected PII entities
        """
        entities = []

        # Phase 1: Regex patterns (fast, high precision)
        regex_entities = self._detect_regex(text)
        entities.extend(regex_entities)

        # Phase 2: Heuristic name detection
        name_entities = self._detect_names_heuristic(text)
        entities.extend(name_entities)

        # Phase 3: LLM analysis (single pass with smart sampling)
        if self.use_llm:
            llm_entities = self._detect_llm_single_pass(text, existing=entities)
            entities.extend(llm_entities)

        # Dedupe and sort
        entities = self._dedupe_entities(entities)
        entities.sort(key=lambda e: e.start if e.start >= 0 else 999999)

        logger.info(f"PII Detection: {len(entities)} entities found")
        return entities

    def _detect_regex(self, text: str) -> List[PIIEntity]:
        """Fast regex-based detection"""
        entities = []

        for pii_type, pattern in self.PATTERNS.items():
            for match in re.finditer(pattern, text, re.IGNORECASE):
                entities.append(
                    PIIEntity(
                        text=match.group(),
                        pii_type=pii_type,
                        start=match.start(),
                        end=match.end(),
                        confidence=0.95,
                    )
                )

        return entities

    def _detect_names_heuristic(self, text: str) -> List[PIIEntity]:
        """
        Detect likely person names using heuristics.

        Patterns detected:
        - "FirstName LastName" (two consecutive capitalized words)
        - Title + Name: "Mr. Smith", "Dr. Jane Doe", "P.Eng. John"
        - Name with suffix: "John Smith, P.Eng.", "Jane Doe, M.Sc."
        - Name in context: "prepared by John Smith", "contact Jane Doe"
        """
        entities = []

        # Pattern: Title + Name (Mr., Mrs., Ms., Dr., Eng., etc.)
        title_pattern = (
            r"\b(?:Mr\.?|Mrs\.?|Ms\.?|Dr\.?|Prof\.?|Eng\.?|P\.?Eng\.?|M\.?Sc\.?)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b"
        )
        for match in re.finditer(title_pattern, text):
            full_match = match.group(0)
            entities.append(
                PIIEntity(
                    text=full_match,
                    pii_type=PIIType.PERSON_NAME,
                    start=match.start(),
                    end=match.end(),
                    confidence=0.9,
                )
            )

        # Pattern: Two capitalized words (FirstName LastName)
        # Must not be preceded by common contextual words
        name_pattern = r"\b([A-Z][a-z]{1,15})\s+([A-Z][a-z]{1,15})\b"
        for match in re.finditer(name_pattern, text):
            first, last = match.group(1), match.group(2)

            # Skip if either part is a stopword
            if first.lower() in self.NAME_STOPWORDS or last.lower() in self.NAME_STOPWORDS:
                continue

            # Skip if looks like a place or org (e.g., "New York", "City Hall")
            if first.lower() in {"new", "old", "city", "lake", "river", "mount"}:
                continue

            full_name = match.group(0)
            entities.append(
                PIIEntity(
                    text=full_name,
                    pii_type=PIIType.PERSON_NAME,
                    start=match.start(),
                    end=match.end(),
                    confidence=0.7,  # Lower confidence for pattern-only
                )
            )

        # Pattern: Name with professional suffix
        suffix_pattern = (
            r"\b([A-Z][a-z]+\s+[A-Z][a-z]+),?\s*(?:P\.?Eng\.?|M\.?Sc\.?|Ph\.?D\.?|B\.?Sc\.?|P\.?Geo\.?|CET)\b"
        )
        for match in re.finditer(suffix_pattern, text):
            name_part = match.group(1)
            entities.append(
                PIIEntity(
                    text=match.group(0),
                    pii_type=PIIType.PERSON_NAME,
                    start=match.start(),
                    end=match.end(),
                    confidence=0.9,
                )
            )

        # Pattern: Contextual names (preceded by trigger words)
        triggers = [
            "prepared by",
            "reviewed by",
            "authored by",
            "contact",
            "attention",
            "attn",
            "signed",
            "approved by",
            "submitted by",
            "client:",
            "owner:",
            "property owner",
            "attention of",
            "care of",
            "c/o",
        ]
        for trigger in triggers:
            pattern = rf"\b{re.escape(trigger)}\s*:?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b"
            for match in re.finditer(pattern, text, re.IGNORECASE):
                name = match.group(1)
                if name.lower() not in self.NAME_STOPWORDS:
                    entities.append(
                        PIIEntity(
                            text=name,
                            pii_type=PIIType.PERSON_NAME,
                            start=match.start(1),
                            end=match.end(1),
                            confidence=0.85,
                        )
                    )

        return entities

    def _detect_llm_single_pass(self, text: str, existing: List[PIIEntity]) -> List[PIIEntity]:
        """
        Single-pass LLM detection with smart sampling.

        Most PII appears at the beginning (header, intro with names/companies)
        and end (signatures, contact info, footer) of documents. This samples
        strategically to reduce from 5+ LLM calls to just 1.
        """
        if not self.client:
            return []

        # Smart sampling: header (names, companies) + footer (signatures, contacts)
        header = text[:3000]
        footer = text[-1500:] if len(text) > 4500 else ""

        # Avoid duplicate content if document is short
        if len(text) <= 4500:
            sample = text
        else:
            sample = header + "\n\n[...middle content omitted...]\n\n" + footer

        existing_texts = {e.text.lower() for e in existing}
        skip_list = list(existing_texts)[:20]

        prompt = f"""Find PII to redact before sending to external API.

Extract ONLY:
- Person names (first/last names of people)
- Company/organization names
- Street addresses
- Project/file numbers
- Phone numbers, emails

DO NOT extract:
- Measurements (depths, distances, weights like "3.5 m", "10 ft", "5 kg")
- Technical values, percentages, coordinates
- Generic terms or engineering terminology

Already found (skip): {skip_list}

Text:
---
{sample}
---

Return JSON array: [{{"text": "exact text", "type": "person_name|company_name|address|project_number|phone|email|location"}}]
Return [] if none found."""

        try:
            logger.info("LLM analyzing document (single-pass)")
            response = self.client.chat(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                options={"temperature": 0.1, "max_tokens": 2048},
            )

            content = response["message"]["content"]

            # Extract JSON from response (handle markdown code blocks)
            content = re.sub(r"```json\s*", "", content)
            content = re.sub(r"```\s*", "", content)

            # Try direct JSON parse first (handles {"entities": [...]} wrapper)
            items = None
            try:
                parsed = json.loads(content)
                if isinstance(parsed, list):
                    items = parsed
                elif isinstance(parsed, dict):
                    for v in parsed.values():
                        if isinstance(v, list):
                            items = v
                            break
            except json.JSONDecodeError:
                # Fall back to regex extraction (non-greedy)
                json_match = re.search(r"\[.*?\]", content, re.DOTALL)
                if json_match:
                    items = json.loads(json_match.group())

            if items is not None:
                entities = []
                for item in items:
                    if isinstance(item, dict) and "text" in item:
                        entity_text = item["text"].strip()

                        # Skip if already found or too short
                        if entity_text.lower() in existing_texts or len(entity_text) < 2:
                            continue

                        # Skip measurements (e.g., "3.5 m", "10 ft", "5.0 kg")
                        if self._is_measurement(entity_text):
                            continue

                        pii_type = self._parse_pii_type(item.get("type", "custom"))
                        entities.append(
                            PIIEntity(
                                text=entity_text,
                                pii_type=pii_type,
                                confidence=0.85,
                            )
                        )

                return entities

        except json.JSONDecodeError as e:
            logger.warning(f"LLM returned invalid JSON: {e}")
        except Exception as e:
            logger.warning(f"LLM detection failed: {e}")

        return []

    def _is_measurement(self, text: str) -> bool:
        """Check if text looks like a measurement (depth, distance, weight, etc.)"""
        # Pattern: number (with optional decimal) followed by unit
        # e.g., "3.5 m", "10 ft", "5.0 kg", "2.5m", "100%"
        measurement_pattern = r"^\s*-?\d+\.?\d*\s*(?:m|cm|mm|km|ft|in|kg|g|mg|lb|oz|%|°|deg|psi|kPa|MPa)\s*$"
        if re.match(measurement_pattern, text, re.IGNORECASE):
            return True

        # Pattern: just a number or decimal
        if re.match(r"^\s*-?\d+\.?\d*\s*$", text):
            return True

        # Pattern: depth ranges like "0.0-3.5 m" or "3.5 - 5.0 m"
        range_pattern = r"^\s*-?\d+\.?\d*\s*[-–—to]+\s*-?\d+\.?\d*\s*(?:m|cm|mm|km|ft|in)?\s*$"
        if re.match(range_pattern, text, re.IGNORECASE):
            return True

        return False

    def _parse_pii_type(self, type_str: str) -> PIIType:
        """Parse PII type string to enum"""
        type_map = {
            "person_name": PIIType.PERSON_NAME,
            "company_name": PIIType.COMPANY_NAME,
            "address": PIIType.ADDRESS,
            "phone": PIIType.PHONE,
            "email": PIIType.EMAIL,
            "project_number": PIIType.PROJECT_NUMBER,
            "file_number": PIIType.FILE_NUMBER,
            "date": PIIType.DATE,
            "location": PIIType.LOCATION,
            "coordinates": PIIType.COORDINATES,
            "license_plate": PIIType.LICENSE_PLATE,
            "sin_ssn": PIIType.SIN_SSN,
            "financial": PIIType.FINANCIAL,
        }
        return type_map.get(type_str.lower(), PIIType.CUSTOM)

    def _dedupe_entities(self, entities: List[PIIEntity]) -> List[PIIEntity]:
        """Remove duplicate entities, keeping highest confidence"""
        seen = {}

        for entity in entities:
            key = entity.text.lower()
            if key not in seen or entity.confidence > seen[key].confidence:
                seen[key] = entity

        return list(seen.values())


class WordRedactor:
    """Redact PII from Word documents"""

    def __init__(self):
        pass

    def redact(
        self,
        input_path: Path,
        output_path: Path,
        entities: List[PIIEntity],
    ) -> Tuple[int, List[str]]:
        """
        Redact entities from Word document.

        Args:
            input_path: Source document
            output_path: Output path for redacted document
            entities: Entities to redact

        Returns:
            Tuple of (redaction_count, warnings)
        """
        from docx import Document

        doc = Document(str(input_path))
        redaction_count = 0
        warnings = []

        # Build replacement map (case-insensitive)
        replacements = {e.text.lower(): e.replacement for e in entities}

        # Process paragraphs
        for para in doc.paragraphs:
            new_text, count = self._replace_in_text(para.text, replacements)
            if count > 0:
                # Preserve formatting by replacing run text
                for run in para.runs:
                    run_new, run_count = self._replace_in_text(run.text, replacements)
                    if run_count > 0:
                        run.text = run_new
                        redaction_count += run_count

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run_new, run_count = self._replace_in_text(run.text, replacements)
                            if run_count > 0:
                                run.text = run_new
                                redaction_count += run_count

        # Process headers/footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        for run in para.runs:
                            run_new, run_count = self._replace_in_text(run.text, replacements)
                            if run_count > 0:
                                run.text = run_new
                                redaction_count += run_count

            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        for run in para.runs:
                            run_new, run_count = self._replace_in_text(run.text, replacements)
                            if run_count > 0:
                                run.text = run_new
                                redaction_count += run_count

        doc.save(str(output_path))
        return redaction_count, warnings

    def _replace_in_text(self, text: str, replacements: Dict[str, str]) -> Tuple[str, int]:
        """Replace all occurrences, case-insensitive"""
        count = 0
        result = text

        for original, replacement in replacements.items():
            pattern = re.compile(re.escape(original), re.IGNORECASE)
            new_result, n = pattern.subn(replacement, result)
            if n > 0:
                result = new_result
                count += n

        return result, count


class PDFRedactor:
    """Redact PII from PDF documents"""

    def __init__(self):
        pass

    def redact(
        self,
        input_path: Path,
        output_path: Path,
        entities: List[PIIEntity],
    ) -> Tuple[int, List[str]]:
        """
        Redact entities from PDF.

        Uses PyMuPDF's redaction feature for proper redaction
        (removes underlying text, not just visual overlay).

        Args:
            input_path: Source PDF
            output_path: Output path
            entities: Entities to redact

        Returns:
            Tuple of (redaction_count, warnings)
        """
        import fitz

        doc = fitz.open(str(input_path))
        redaction_count = 0
        warnings = []

        # Search for each entity across all pages
        for entity in entities:
            search_text = entity.text

            for page_num, page in enumerate(doc):
                # Search for text instances (case-insensitive: try variations)
                text_instances = page.search_for(search_text)
                if not text_instances:
                    text_instances = page.search_for(search_text.lower())
                if not text_instances:
                    text_instances = page.search_for(search_text.upper())
                if not text_instances:
                    text_instances = page.search_for(search_text.title())

                for inst in text_instances:
                    # Add redaction annotation
                    page.add_redact_annot(
                        inst,
                        text=entity.replacement,
                        fontsize=10,
                        fill=(1, 1, 1),  # White background
                    )
                    redaction_count += 1

        # Apply all redactions
        for page in doc:
            page.apply_redactions()

        doc.save(str(output_path))
        doc.close()

        return redaction_count, warnings


class Redactrr:
    """
    Main redaction pipeline.

    Usage:
        redactor = Redactrr()
        result = redactor.redact("/path/to/document.docx")
        print(f"Redacted: {result.redacted_path}")
    """

    SUPPORTED_EXTENSIONS = {".docx", ".pdf"}

    def __init__(self, use_llm: bool = True, model: str = None):
        """
        Args:
            use_llm: Whether to use LLM for context-sensitive detection
            model: Model name for LLM detection
        """
        self.detector = PIIDetector(use_llm=use_llm, model=model)
        self.word_redactor = WordRedactor()
        self.pdf_redactor = PDFRedactor()

    def redact(
        self,
        input_path: Path,
        output_path: Path = None,
        additional_terms: List[str] = None,
    ) -> RedactionResult:
        """
        Redact PII from document.

        Args:
            input_path: Source document
            output_path: Output path (default: input_REDACTED.ext)
            additional_terms: Extra terms to redact

        Returns:
            RedactionResult with status and statistics
        """
        input_path = Path(input_path)
        warnings = []

        # Validate
        if not input_path.exists():
            return RedactionResult(
                original_path=str(input_path),
                redacted_path="",
                success=False,
                entities_found=0,
                entities_redacted=0,
                entity_types={},
                error=f"File not found: {input_path}",
            )

        suffix = input_path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            return RedactionResult(
                original_path=str(input_path),
                redacted_path="",
                success=False,
                entities_found=0,
                entities_redacted=0,
                entity_types={},
                error=f"Unsupported format: {suffix}. Supported: {self.SUPPORTED_EXTENSIONS}",
            )

        # Default output path
        if output_path is None:
            output_path = input_path.parent / f"{input_path.stem}_REDACTED{suffix}"
        output_path = Path(output_path)

        logger.info(f"Redacting: {input_path.name}")

        try:
            # Extract text for analysis
            text = self._extract_text(input_path)

            if not text:
                return RedactionResult(
                    original_path=str(input_path),
                    redacted_path="",
                    success=False,
                    entities_found=0,
                    entities_redacted=0,
                    entity_types={},
                    error="PDF appears to be scanned/image-only (no extractable text)",
                )

            # Detect PII
            entities = self.detector.detect(text)

            # Add custom terms and deduplicate
            if additional_terms:
                for term in additional_terms:
                    entities.append(
                        PIIEntity(
                            text=term,
                            pii_type=PIIType.CUSTOM,
                        )
                    )
                entities = self.detector._dedupe_entities(entities)

            if not entities:
                # No PII found - just copy file
                import shutil

                shutil.copy(input_path, output_path)

                return RedactionResult(
                    original_path=str(input_path),
                    redacted_path=str(output_path),
                    success=True,
                    entities_found=0,
                    entities_redacted=0,
                    entity_types={},
                    warnings=["No PII detected - file copied without changes"],
                )

            # Apply redactions
            if suffix == ".docx":
                redaction_count, redact_warnings = self.word_redactor.redact(input_path, output_path, entities)
            else:  # .pdf
                redaction_count, redact_warnings = self.pdf_redactor.redact(input_path, output_path, entities)

            warnings.extend(redact_warnings)

            # Build entity type counts
            type_counts = {}
            for entity in entities:
                type_name = entity.pii_type.value
                type_counts[type_name] = type_counts.get(type_name, 0) + 1

            logger.info(f"Redacted {redaction_count} instances of {len(entities)} entities")

            return RedactionResult(
                original_path=str(input_path),
                redacted_path=str(output_path),
                success=True,
                entities_found=len(entities),
                entities_redacted=redaction_count,
                entity_types=type_counts,
                warnings=warnings,
            )

        except Exception as e:
            error_msg = str(e)
            if "encrypt" in error_msg.lower() or "password" in error_msg.lower():
                error_msg = "PDF appears to be encrypted"
            else:
                error_msg = f"Failed to extract text from document: {error_msg}"
            logger.error(f"Redaction failed: {error_msg}")
            return RedactionResult(
                original_path=str(input_path),
                redacted_path="",
                success=False,
                entities_found=0,
                entities_redacted=0,
                entity_types={},
                error=error_msg,
            )

    def _extract_text(self, file_path: Path) -> str:
        """Extract text from document for analysis"""
        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            from docx import Document

            doc = Document(str(file_path))

            parts = []
            for para in doc.paragraphs:
                parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)

            return "\n".join(parts)

        elif suffix == ".pdf":
            import fitz

            doc = fitz.open(str(file_path))

            parts = []
            for page in doc:
                parts.append(page.get_text())
            doc.close()

            return "\n".join(parts)

        return ""

    def preview(self, input_path: Path) -> Dict[str, Any]:
        """
        Preview what would be redacted without making changes.

        Args:
            input_path: Document to analyze

        Returns:
            Dict with detected entities and their types
        """
        input_path = Path(input_path)

        text = self._extract_text(input_path)
        if not text:
            return {"error": "Could not extract text", "entities": []}

        entities = self.detector.detect(text)

        return {
            "file": str(input_path),
            "text_length": len(text),
            "entities_found": len(entities),
            "entities": [
                {
                    "text": e.text,
                    "type": e.pii_type.value,
                    "replacement": e.replacement,
                    "confidence": e.confidence,
                }
                for e in entities
            ],
            "by_type": self._count_by_type(entities),
        }

    def _count_by_type(self, entities: List[PIIEntity]) -> Dict[str, int]:
        """Count entities by type"""
        counts = {}
        for e in entities:
            type_name = e.pii_type.value
            counts[type_name] = counts.get(type_name, 0) + 1
        return counts


def redact_document(
    input_path: Path,
    output_path: Path = None,
    use_llm: bool = True,
    additional_terms: List[str] = None,
) -> RedactionResult:
    """
    Convenience function to redact a document.

    Args:
        input_path: Source document
        output_path: Output path (default: auto-generated)
        use_llm: Whether to use LLM for detection
        additional_terms: Extra terms to redact

    Returns:
        RedactionResult
    """
    redactor = Redactrr(use_llm=use_llm)
    return redactor.redact(input_path, output_path, additional_terms)


def preview_redaction(input_path: Path, use_llm: bool = True) -> Dict[str, Any]:
    """Preview what would be redacted"""
    redactor = Redactrr(use_llm=use_llm)
    return redactor.preview(input_path)
